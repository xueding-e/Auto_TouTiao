#!/usr/bin/env python3
"""
鲲写作 Markdown → PDF / 头条富文本 / Docx 转换脚本
用法:
  python md_to_pdf.py input.md output.docx --docx                   # 输出 docx（推荐！头条可直接导入，格式完整）
  python md_to_pdf.py input.md output.html --toutiao                # 输出干净 HTML（可复制粘贴到头条）
  python md_to_pdf.py input.md output.pdf [--title "标题"]          # 输出 PDF 存档

依赖:
  pip install markdown --break-system-packages                      # 基础依赖
  pip install python-docx --break-system-packages                   # docx 输出（推荐）
  如需 PDF: pip install playwright && playwright install chromium   # 或 pip install weasyprint
"""

import sys
import os
import re
import argparse
import markdown
import tempfile

# 导入统一配置加载器
_project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
if _project_root not in sys.path:
    sys.path.insert(0, _project_root)
from config.config_loader import get_user_agent, get_timeout
import hashlib
from html.parser import HTMLParser
from pathlib import Path
from io import BytesIO


# ── 封面图生成 ──

def generate_cover_image(title, output_path, width=800, height=450):
    """用 Pillow 生成带标题的封面图"""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("[WARN] Pillow 未安装，跳过封面图生成。pip install Pillow --break-system-packages")
        return None

    img = Image.new('RGB', (width, height), color=(26, 82, 118))  # 深蓝底
    draw = ImageDraw.Draw(img)

    # 尝试加载中文字体
    font_paths = [
        'C:/Windows/Fonts/msyh.ttc',      # 微软雅黑
        'C:/Windows/Fonts/simhei.ttf',     # 黑体
        'C:/Windows/Fonts/simsun.ttc',     # 宋体
        '/System/Library/Fonts/PingFang.ttc',  # Mac
        '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',  # Linux
    ]
    title_font = None
    sub_font = None
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                title_font = ImageFont.truetype(fp, 38)
                sub_font = ImageFont.truetype(fp, 20)
                break
            except Exception:
                continue
    if title_font is None:
        title_font = ImageFont.load_default()
        sub_font = ImageFont.load_default()

    # 排版标题（自动换行）
    max_line_width = width - 120
    chars_per_line = 16  # 每行约16个字
    title_lines = []
    for k in range(0, len(title), chars_per_line):
        title_lines.append(title[k:k+chars_per_line])

    # 绘制装饰线
    line_y = height // 2 - 30 * len(title_lines) // 2 - 20
    draw.rectangle([60, line_y - 5, 100, line_y - 3], fill=(46, 134, 193))  # 蓝色装饰条

    # 绘制标题
    for j, tl in enumerate(title_lines):
        tl_bbox = draw.textbbox((0, 0), tl, font=title_font)
        tl_w = tl_bbox[2] - tl_bbox[0]
        tl_x = (width - tl_w) // 2
        draw.text((tl_x, line_y + j * 52), tl, fill=(255, 255, 255), font=title_font)

    # 底部作者标签
    footer_y = height - 65
    draw.rectangle([width//2 - 80, footer_y, width//2 + 80, footer_y + 28], fill=(46, 134, 193))
    author_txt = '数字生命鲲'
    auth_bbox = draw.textbbox((0, 0), author_txt, font=sub_font)
    auth_w = auth_bbox[2] - auth_bbox[0]
    draw.text(((width - auth_w) // 2, footer_y + 3), author_txt, fill=(255, 255, 255), font=sub_font)

    img.save(output_path, quality=90)
    return output_path


# ── 图片处理 ──

def _fetch_image(image_path_or_url):
    """获取图片内容：本地文件 → 直接读取；URL → 下载；都不行 → None"""
    # 本地文件
    if os.path.exists(image_path_or_url):
        with open(image_path_or_url, 'rb') as f:
            return BytesIO(f.read())
    # URL
    if image_path_or_url.startswith(('http://', 'https://')):
        try:
            import requests
            headers = {
                'User-Agent': get_user_agent(),
                'Referer': 'https://image.baidu.com/',
            }
            resp = requests.get(image_path_or_url, headers=headers, timeout=get_timeout('default'))
            resp.raise_for_status()
            return BytesIO(resp.content)
        except Exception:
            pass
    return None


def _insert_image_placeholder(doc, description):
    """插入一个带描述的图片占位框"""
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 占位段落
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)

    # 灰色占位框（用表格模拟）
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.width = Cm(15)

    # 设置单元格高度
    cell_p = cell.paragraphs[0]
    cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_p.paragraph_format.space_before = Pt(60)
    cell_p.paragraph_format.space_after = Pt(60)

    # 图片图标 + 描述
    run_icon = cell_p.add_run('🖼️ ')
    run_icon.font.size = Pt(24)
    run_desc = cell_p.add_run(description if description else '配图')
    run_desc.font.size = Pt(11)
    run_desc.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    run_desc.italic = True

    # 在表格下方加一个空行做间距
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return cell_p


def _insert_image(doc, image_data, description='', max_width_cm=15):
    """插入图片到docx"""
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)

    try:
        run = para.add_run()
        run.add_picture(image_data, width=Cm(max_width_cm))
    except Exception:
        return None

    # 图片描述（如果有）
    if description:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        run_cap = cap.add_run(description)
        run_cap.font.size = Pt(9)
        from docx.shared import RGBColor
        run_cap.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
        run_cap.italic = True

    return para


# ── Markdown → Docx 转换 ──

def _add_formatted_paragraph(doc, text, style='Normal', bold_pattern=None):
    """添加一个段落，支持 **粗体** 内联格式"""
    para = doc.add_paragraph(style=style)
    if not text:
        return para

    # 按 **...** 分割文本，交替普通/粗体
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:  # 奇数索引 = 粗体内容
            run.bold = True
    return para


def md_to_docx(md_text, output_path):
    """将 Markdown 转为 .docx 文件（头条可直接导入）"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[ERROR] python-docx 未安装。请执行: pip install python-docx --break-system-packages")
        raise

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.6

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # H1 标题 → docx Heading 1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:]
            _add_formatted_paragraph(doc, title_text, style='Heading 1')
            i += 1
            continue

        # H2 标题 → docx Heading 2
        if stripped.startswith('## '):
            _add_formatted_paragraph(doc, stripped[3:], style='Heading 2')
            i += 1
            continue

        # 引用块（> 开头）
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            para = _add_formatted_paragraph(doc, quote_text, style='Quote')
            # 如果连续多行都是引用，合并
            i += 1
            while i < len(lines) and lines[i].strip().startswith('> '):
                run = para.add_run('\n' + lines[i].strip()[2:])
                i += 1
            continue

        # 分隔线
        if stripped in ('---', '***', '___'):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            # 添加一条横线（用下划线字符模拟）
            run = para.add_run('─' * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # 图片：![描述](路径或URL)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            desc = img_match.group(1)
            src = img_match.group(2)
            img_data = _fetch_image(src)
            if img_data:
                _insert_image(doc, img_data, desc)
            else:
                _insert_image_placeholder(doc, desc or src)
            i += 1
            continue

        # 普通段落：合并连续非空行
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('# ', '> ', '---', '***', '___')):
            para_lines.append(lines[i].strip())
            i += 1
        full_text = ' '.join(para_lines)
        _add_formatted_paragraph(doc, full_text)

    doc.save(output_path)
    return True


# ── 头条 HTML 清洗 ──


# ── 头条编辑器支持的标签白名单 ──
TOUTIAO_ALLOWED_TAGS = {'p', 'strong', 'b', 'em', 'i', 'h2', 'h3', 'h4',
                        'blockquote', 'br', 'hr', 'ul', 'ol', 'li', 'a', 'span'}


class ToutiaoHTMLCleaner(HTMLParser):
    """清洗 HTML，只保留头条编辑器支持的标签和内容"""
    def __init__(self):
        super().__init__()
        self.result = []
        self.skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if self.skip_depth > 0:
            self.skip_depth += 1
            return
        if tag in TOUTIAO_ALLOWED_TAGS:
            # 只保留标签本身，去掉所有属性（头条会自己加样式）
            if tag == 'a':
                # 保留链接的 href
                href = dict(attrs).get('href', '')
                if href:
                    self.result.append(f'<a href="{href}">')
                else:
                    self.result.append(f'<{tag}>')
            else:
                self.result.append(f'<{tag}>')
        elif tag in {'h1', 'title'}:
            # h1 转为 h2（头条编辑器用 h2 做主标题）
            self.result.append('<h2>')
        elif tag in {'style', 'script', 'div', 'section', 'article', 'header', 'footer',
                      'nav', 'aside', 'main', 'figure', 'figcaption', 'pre', 'code'}:
            self.skip_depth = 1

    def handle_endtag(self, tag):
        if self.skip_depth > 0:
            self.skip_depth -= 1
            return
        if tag in TOUTIAO_ALLOWED_TAGS:
            self.result.append(f'</{tag}>')
        elif tag in {'h1', 'title'}:
            self.result.append('</h2>')
        elif tag in {'style', 'script', 'div', 'section', 'article', 'header', 'footer',
                      'nav', 'aside', 'main', 'figure', 'figcaption', 'pre', 'code'}:
            pass  # 跳过结束标签

    def handle_data(self, data):
        if self.skip_depth > 0:
            return
        self.result.append(data)

    def handle_startendtag(self, tag, attrs):
        if tag in {'br', 'hr'}:
            self.result.append(f'<{tag}>')

    def get_clean_html(self):
        return ''.join(self.result)


def clean_for_toutiao(html_string):
    """将 Markdown 转出的 HTML 清洗为头条编辑器兼容格式"""
    # 1. 用 HTMLParser 剥离 CSS / div / 非白名单标签
    cleaner = ToutiaoHTMLCleaner()
    cleaner.feed(html_string)
    clean = cleaner.get_clean_html()

    # 2. 修复空段落：多个连续 <br> 转为一个 <p>
    clean = re.sub(r'(<br>\s*){2,}', '</p><p>', clean)

    # 3. 去掉开头的空标签
    clean = re.sub(r'^(<br>|<p></p>|</p>)+', '', clean)

    # 4. 去掉末尾的空标签
    clean = re.sub(r'(<br>|<p></p>|</p>)+$', '', clean)

    # 5. 确保以 <p> 开头（如果不是标题标签的话）
    if not re.match(r'^\s*<[hb]', clean):
        clean = '<p>' + clean
    # 确保以 </p> 结尾
    if not re.search(r'</[hp]>\s*$', clean):
        clean = clean + '</p>'

    # 6. 压缩多余空白（保留标签间的换行便于阅读）
    clean = re.sub(r'\n\s*\n', '\n', clean)
    clean = clean.strip()

    return clean


def md_to_toutiao_html(md_text):
    """将 Markdown 转为头条编辑器可直接粘贴的干净 HTML"""
    # 用 markdown 库转换
    html_body = markdown.markdown(
        md_text,
        extensions=['tables', 'fenced_code'],
        output_format='html5'
    )
    # 清洗为头条兼容格式
    return clean_for_toutiao(html_body)


# ── CSS 样式（文章阅读优化） ──
CSS_TEMPLATE = """
@page {
    size: A4;
    margin: 20mm 25mm 20mm 25mm;

    @bottom-center {
        content: "第 " counter(page) " 页";
        font-family: "Droid Sans Fallback", Helvetica, Arial, sans-serif;
        font-size: 8pt;
        color: #95a5a6;
        border-top: 0.5pt solid #d5dbdb;
        padding-top: 2mm;
    }
}

@page :first {
    @bottom-center { content: none; }
}

body {
    font-family: "Droid Sans Fallback", Helvetica, Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.8;
    color: #2c3e50;
    text-align: justify;
}

/* 封面 */
.cover {
    page-break-after: always;
    text-align: center;
    padding-top: 40%;
}
.cover h1 {
    font-size: 24pt;
    color: #1a5276;
    margin-bottom: 10mm;
    font-weight: bold;
    letter-spacing: 1.5pt;
    line-height: 1.4;
}
.cover .meta {
    font-size: 11pt;
    color: #7f8c8d;
    margin-bottom: 4mm;
}
.cover .divider {
    width: 50%;
    margin: 10mm auto;
    border: none;
    border-top: 1pt solid #1a5276;
}

/* 一级标题（文章主标题） */
h1 {
    font-size: 20pt;
    color: #1a5276;
    margin-top: 12mm;
    margin-bottom: 6mm;
    padding-bottom: 2mm;
    border-bottom: 1.5pt solid #1a5276;
    page-break-before: always;
    font-weight: bold;
}

/* 二级标题 */
h2 {
    font-size: 14pt;
    color: #1e8449;
    margin-top: 8mm;
    margin-bottom: 4mm;
    font-weight: bold;
}

/* 三级标题 */
h3 {
    font-size: 12pt;
    color: #2e86c1;
    margin-top: 6mm;
    margin-bottom: 3mm;
    font-weight: bold;
}

/* 段落 */
p {
    margin-top: 2mm;
    margin-bottom: 2mm;
    orphans: 2;
    widows: 2;
}

/* 引用块 */
blockquote {
    margin: 4mm 0;
    padding: 3mm 4mm 3mm 8mm;
    background: #f8f9fa;
    border-left: 3pt solid #1a5276;
    color: #5d6d7e;
    font-size: 10pt;
}
blockquote p {
    margin: 1mm 0;
}

/* 粗体 / 金句 */
strong, b {
    font-weight: bold;
    color: #1a252f;
}

/* 行内代码 */
code {
    font-family: "Courier New", Courier, monospace;
    background: #fdf2e9;
    color: #c0392b;
    padding: 0.5mm 1.5mm;
    border-radius: 2pt;
    font-size: 9.5pt;
}

/* 分隔线 */
hr {
    border: none;
    border-top: 0.5pt solid #bdc3c7;
    margin: 5mm 0;
}

/* 列表 */
ul, ol {
    margin: 2mm 0;
    padding-left: 8mm;
}
li {
    margin-bottom: 1.5mm;
}

/* 链接 */
a {
    color: #2e86c1;
    text-decoration: none;
}
"""


def md_to_html(md_text, title="文章标题", subtitle="", meta_line="", author="鲲"):
    """将 Markdown 转为带封面的 HTML"""

    # 用 markdown 库转换正文
    html_body = markdown.markdown(
        md_text,
        extensions=['tables', 'fenced_code', 'nl2br'],
        output_format='html5'
    )

    # 移除正文中的第一个 h1（会用在封面上）
    first_h1_match = re.search(r'<h1>(.*?)</h1>', html_body)
    if first_h1_match:
        extracted_title = first_h1_match.group(1)
        if not title or title == "文章标题":
            title = extracted_title
        html_body = html_body.replace(first_h1_match.group(0), '', 1)

    # 构建封面
    cover_html = f"""
    <div class="cover">
        <h1 style="page-break-before: avoid; border: none;">{title}</h1>
        {"<div class='meta'>" + subtitle + "</div>" if subtitle else ""}
        {"<div class='meta'>" + meta_line + "</div>" if meta_line else ""}
        <hr class="divider">
        <div class="meta">作者: {author}</div>
    </div>
    """

    full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <style>{CSS_TEMPLATE}</style>
</head>
<body>
{cover_html}
{html_body}
</body>
</html>"""

    return full_html


def html_to_pdf_playwright(html_string, output_path):
    """使用 Playwright 将 HTML 转为 PDF"""
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_string, wait_until='networkidle')
        page.pdf(
            path=output_path,
            format='A4',
            margin={'top': '20mm', 'bottom': '20mm', 'left': '25mm', 'right': '25mm'},
            print_background=True
        )
        browser.close()
    return True


def html_to_pdf_weasyprint(html_string, output_path):
    """使用 WeasyPrint 将 HTML 转为 PDF"""
    from weasyprint import HTML
    HTML(string=html_string).write_pdf(output_path)
    return True


def html_to_pdf(html_string, output_path):
    """智能选择 PDF 生成引擎"""
    # 优先尝试 WeasyPrint（排版更好），失败则回退到 Playwright
    try:
        from weasyprint import HTML
        HTML(string=html_string).write_pdf(output_path)
        print("[OK] 使用 WeasyPrint 引擎生成 PDF")
        return
    except (ImportError, OSError) as e:
        print(f"[INFO] WeasyPrint 不可用 ({type(e).__name__})，尝试 Playwright...")

    try:
        html_to_pdf_playwright(html_string, output_path)
        print("[OK] 使用 Playwright 引擎生成 PDF")
        return
    except ImportError:
        print("[ERROR] Playwright 未安装。请执行: pip install playwright && playwright install chromium")
        raise
    except Exception as e:
        print(f"[ERROR] Playwright 生成 PDF 失败: {e}")
        raise


def main():
    parser = argparse.ArgumentParser(description="鲲写作 Markdown → PDF / 头条富文本 / Docx")
    parser.add_argument("input", help="输入的 Markdown 文件路径")
    parser.add_argument("output", help="输出的 PDF / HTML / Docx 文件路径")
    parser.add_argument("--title", default=None, help="文章标题（PDF模式用，默认从 Markdown 第一个 H1 提取）")
    parser.add_argument("--author", default="鲲", help="作者名")
    parser.add_argument("--toutiao", action="store_true",
                        help="输出头条编辑器兼容的干净 HTML（直接粘贴即可保留格式）")
    parser.add_argument("--docx", action="store_true",
                        help="输出 .docx 文档（推荐！头条可直接导入，粗体/段落/标题完整保留）")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        md_text = f.read()

    # ── Docx 模式：生成 .docx 文件（头条推荐导入格式）──
    if args.docx:
        md_to_docx(md_text, args.output)
        size_kb = os.path.getsize(args.output) / 1024
        print(f"[OK] Docx 已生成: {args.output} ({size_kb:.1f} KB)")
        print(f"      可直接在头条编辑器「导入文档」上传此文件，格式完整保留")
        return

    # ── 头条 HTML 模式 ──
    if args.toutiao:
        toutiao_html = md_to_toutiao_html(md_text)
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(toutiao_html)
        print(f"[OK] 头条富文本已生成: {args.output}")
        print(f"      可直接复制内容粘贴到头条编辑器，粗体/段落格式完整保留")
        return

    # ── PDF 模式 ──

    # 提取元信息（作者/日期行）
    meta_line = ""
    for line in md_text.split("\n"):
        stripped = line.strip().lstrip(">").strip()
        if "作者" in stripped or "日期" in stripped:
            meta_line = stripped
            break

    subtitle = ""
    html = md_to_html(md_text, title=args.title or "文章标题",
                      subtitle=subtitle, meta_line=meta_line, author=args.author)

    # 保存中间 HTML（便于调试）
    html_path = args.output.replace('.pdf', '.html')
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"[OK] HTML 已生成: {html_path}")

    # 转 PDF
    html_to_pdf(html, args.output)
    size_kb = os.path.getsize(args.output) / 1024
    print(f"[OK] PDF 已生成: {args.output} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
